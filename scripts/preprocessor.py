#!/usr/bin/python
import os
import re
from pathlib import Path

import aspose.pdf as pdf
import repackage
import spacy
from docx import Document

repackage.up()
from config.config import load_config

config = load_config()


def main(path: str | None = None):
    """
    Main preprocessor of different files. Adds chunks of texts to a local CSV file.

    Args:
        path (str | None, optional): Local path files are in. Defaults to None.
    """
    list_of_files = get_files_in_dir(path)
    for file_path in list_of_files:
        if file_path.endswith(".pdf"):
            pass
        elif file_path.endswith(".docx"):
            text = DOCXPreprocessor(file_path)
            list_of_sentences = text.preprocess()
            add_lines(list_of_sentences)
        else:
            continue


def get_files_in_dir(path_to_search: str | None = None) -> list:
    """
    Searches all files in a dir (checking also subdirs) that have extensions as stated
    in config.json.

    Args:
        path_to_search (str | None, optional): Path to search files in. If None,
        equals to self.path. Defaults to None.

    Returns:
        list: List of files with absolute paths.
    """
    if path_to_search is None:
        path_to_search = Path(__file__).parent.parent.joinpath("data")
    if os.path.isfile(path_to_search):
        return [path_to_search]
    elif os.path.isdir(path_to_search):
        file_paths = []
        for root, _, files in os.walk(path_to_search):
            for file in files:
                _, ext = os.path.splitext(file)
                if ext.lower() in config["general"]["accepted_file_formats"]:
                    file_path = os.path.join(root, file)
                    file_paths.append(os.path.abspath(file_path))
        return file_paths


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts sentences from DOCX file and returns as one string. Joins single and
    multiple '\n' into single whitespaces.

    Args:
        file_path (str): DOCX file path to extract text from.

    Returns:s
        str: Extracted text.
    """
    doc = Document(file_path)
    text = " ".join([paragraph.text for paragraph in doc.paragraphs])
    return text


def split_on_points(text: str) -> list[str]:
    """
    Splits text of paragraphs and points (eg. § 2, §2, $ 2, $2, 2., 13.).

    Args:
        text (str): Text to be divided.

    Returns:
        list[str]: List of remained sentences.
    """
    # dzielenie na paragrafach i punktach TODO: czy dzielimy na podpunktach?
    text_without_paragrapghs = re.split(r"[§$] ?\d+.?", text)
    text_without_points = re.split(r"\d+\.", " ".join(text_without_paragrapghs))
    # usuwanie białych znaków na początku i na końcu stringa
    # oraz usuwanie pustych elementów listy
    sentences = [sentence.strip() for sentence in text_without_points if sentence.strip()]
    # usuwanie wielokrotnych spacji wewnątrz stringów i zwrócenie listy
    return [replace_whitespaces(sentence) for sentence in sentences]


def remove_preambule_before_par(text: str) -> str:
    """
    Removes everything before first paragraph (§) including this paragraph sign with a
    number. Strips whitespace at the beginnin of the remained text. As first paragraph
    (§) forms may vary, the following substrings are tried to be splitters: "§ 1",
    "§1", "$ 1", "$1". WARNING: this set may not be a finite one.

    Args:
        text (str): Text to be divided.

    Returns:
        str: Remained text body.
    """
    # usuwanie wstępu (przed paragrafem 1.)
    try:
        return re.split(r"[§$] ?1.?", text, 1)[1]
    except IndexError:
        return text


def remove_preambule_before_point(text: str) -> str:
    """
    Removes everything before first point (1.).

    Args:
        text (str): Text to be divided.

    Returns:
        str: Remained text body.
    """
    # usuwanie wstępu (przed punktem 1.)
    try:
        return "1. " + re.split(r" 1.? ", text, 1)[1]
    except IndexError:
        return text


def remove_attachments(text: str) -> str:
    """
    Removes everything before first attachment (Załącznik nr 1).

    Args:
        text (str): Text to be divided.

    Returns:
        str: Remained text body.
    """
    # usuwanie załączników wraz z tytułem sekcji (np. "Załącznik nr X")
    return re.split(r"Za[lł][aą]cznik ?(nr)? ?1", text, 1)[0]


def replace_forbidden_chars(
    text: str | list[str],
    forbidden_chars: list[str] | None = None,
    replace_with: str = ",",
) -> str | list[str]:
    # TODO: implement dictionary that maps characters to be replaced to characters \
    # to replace with
    """
    Replaces characters that prevent from reading CSV file into pandas DataFrame.

    Args:
        text (str | list[str]): Text or list of texts to apply replacement on.
        forbidden_chars (list[str] | None, optional): Forbidden characters to replace
        with `replace_with`. Defaults to None.
        replace_with (str): Character to replace with. Defaults to ','.
    Returns:
        (str | list[str]): Text or list of texts.
    """
    if forbidden_chars is None:
        forbidden_chars = [";"]
    if isinstance(text, str):
        for char in forbidden_chars:
            text = text.replace(char, replace_with)
        return text
    elif isinstance(text, list):
        result = []
        for elem in text:
            if not isinstance(elem, str):
                raise TypeError("Text must be string or list of strings.")
            for char in forbidden_chars:
                elem = elem.replace(char, replace_with)
            result.append(elem)
        return result
    else:
        raise TypeError("Text must be string or list of strings.")


def replace_whitespaces(text: str | list[str]) -> str:
    """
    Replaces multiple whitespaces with single ones.

    Args:
        text (str | list[str]): Text of list of texts to apply replacement on.

    Returns:
        str | list[str]: Text or list of texts.
    """
    if isinstance(text, str):
        return re.sub(r"\s+", " ", text)
    elif isinstance(text, list):
        return [re.sub(r"\s+", " ", t) for t in text]
    else:
        raise TypeError("Text must be string of list of strings.")


def remove_page_numbers(text: str) -> str:
    """
    Replaces page numbers (i.e. '2 z 5' or '13 z 134').

    Args:
        text (str): Text to apply removal on.

    Returns:
        str: Text.
    """
    return re.sub(r"\d{1,3} z \d{1,3}", "", text)


def add_category(text: str, file_path: str) -> str:
    category = file_path.split("\\")[-1].split(".")[0]
    return category + ": " + text + "\n"


def determine_language(text: str) -> str:
    """
    Returns a language, that given text is written in.

    Args:
        text (str): Text to check.

    Returns:
        str: Language. Tested for pl, en, de.
    """
    import spacy_fastlang

    nlp = spacy.load("pl_core_news_sm")
    nlp.add_pipe("language_detector")
    doc = nlp(text)

    return doc._.language


def chunk_text(text: str) -> list[str]:
    """
    Chunks text into overlapping 3-sentence chunks (they overlap 1 sentence with previous
    chunk, 1 with next and 1 sentence is non-overlapping.)

    Args:
        text (str): String to chunk.

    Returns:
        list[str]: List of strings after chunking.
    """
    nlp = spacy.load("pl_core_news_sm")
    doc = nlp(text)
    sentences = [sent.text.strip() for sent in doc.sents]
    result = []
    for i in range(len(sentences) - 2):
        pattern = " ".join(str(sentences[j]) for j in range(i, i + 3))
        result.append(pattern)
    return result


class DOCXPreprocessor:
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
        if "Zarządzenie" in self.file_path:
            self.file_type = "Zarządzenie"
        elif "Pytania" in self.file_path:
            self.file_type = "Pytania"
        else:
            self.file_type = "Inny"

    def preprocess(self) -> list[str]:
        """
        Sequentially processes DOCX file.

        Returns:
            list[str]: List of sentences
        """
        if self.file_type.capitalize() == "Zarządzenie":
            t1 = extract_text_from_docx(self.file_path)
            t2 = remove_preambule_before_par(t1)
            t3 = remove_attachments(t2)
            t4 = replace_forbidden_chars(t3)
            return split_on_points(t4)
        elif self.file_type.capitalize() == "Pytania":
            t1 = extract_text_from_docx(self.file_path)
            t2 = remove_preambule_before_point(t1)
            t3 = remove_page_numbers(t2)
            t4 = replace_whitespaces(t3)
            t5 = replace_forbidden_chars(t4)
            t6 = add_category(t5, self.file_path)
            return [t6]
        else:
            t1 = extract_text_from_docx(self.file_path)
            t2 = chunk_text(t1)
            t3 = replace_whitespaces(t2)
            return replace_forbidden_chars(t3)


class PDFPreprocessor:
    def get_raw_text_from_pdf(self, file_path: str) -> str:
        """
        Returns raw text (single string) from a PDF file.

        Args:
            file_path (str): File path to read PDF file from.

        Returns:
            str: Full text of a PDF file
        """
        pdfDocument = pdf.Document(file_path)
        full_text = []
        for i in range(0, len(pdfDocument.pages) - 1):
            full_text.append(self.get_raw_text_from_tables(pdfDocument, i + 1))
        return "".join(full_text).replace("  ", " ")

    def get_raw_text_from_tables(self, pdfDocument: pdf.Document, iterator: int) -> str:
        """
        Returns raw text (single string) from a table in a PDF file.

        Args:
            pdfDocument (pdf.Document): PDF pages.
            iterator (int): Number of page.

        Returns:
            str: Full text of a PDF file page.
        """
        # Initialize TableAbsorber object
        tableAbsorber = pdf.text.TableAbsorber()

        # Parse all the tables on first page
        tableAbsorber.visit(pdfDocument.pages[iterator])

        # Get a reference of the first table
        absorbedTable = tableAbsorber.table_list[0]

        full_text = []

        # Iterate through all the rows in the table
        for pdfTableRow in absorbedTable.row_list:
            # Iterate through all the columns in the row
            for pdfTableCell in pdfTableRow.cell_list:
                # Fetch the text fragments
                textFragmentCollection = pdfTableCell.text_fragments
                # Iterate through the text fragments
                for textFragment in textFragmentCollection:
                    # Print the text
                    full_text.extend(textFragment.text)
        return "".join(full_text).replace("  ", " ")

    def jsonize_pdf(self, main_string: str) -> dict:
        """
        Gets saught after elements and return them as a dict, eg.
        dict['Metody kształcenia'] = 'Wskazane przez praktykodawcę'

        Args:
            main_string (str): String to search in.

        Returns:
            dict: Dictionary with `headers` as keys and information as values.
        """
        json_ = {}
        list_of_headers = [
            "Nazwa przedmiotu",
            "Forma zajęć",
            "Rok akademicki, rok studiów, semestr realizacji przedmiotu",
            "Stopień studiów, tryb studiów",
            "Cel przedmiotu",
            "Wymagania wstępne",
            "Metody kształcenia",
        ]
        for elem in list_of_headers:
            json_[elem] = self.forward_regexp_search(main_string, elem)
        return json_


def forward_regexp_search(main_string: str, target_string: str) -> str:
    """
    Searches a given substring (`target_string`) in a large string (`main_string`)
    and returns N next words as one string (up to the nearest number with a comma, eg.
    `13.`)

    Args:
        main_string (str): String to search in
        target_string (str): Subtring to search

    Returns:
        str: Searched string.
    """

    def get_number_with_comma(string):
        regexp = r"\d+\."
        if re.fullmatch(regexp, string) is not None:
            return False
        return True

    # get rid of whitespaces at the beginning and in the end
    target_string = target_string.strip()
    length = len(target_string)
    # position of the target_string
    pos = main_string.find(target_string)
    word_list = (main_string[pos + length :].strip()).split(" ")
    result = []
    for word in word_list:
        if get_number_with_comma(word):
            result.append(word)
        else:
            break
    return " ".join(result)


def add_lines(sentences: list[str], file_path: str | None = None):
    """
    Appends sentences to file. Creates a file if it does no exist.

    Args:
        sentences (list[str]): Sentences to append.
        file_path (str | None, optional): CSV file path to append sentences to.
        If None, taken from config.json. Defaults to None.
    """
    if file_path is None:
        file_name = config["pinecone"]["target_filename"]["raw"]
        file_path = Path(__file__).parent.parent.joinpath(file_name)
    create_csv_file_if_not_exist(file_path)
    with open(file_path, "a", encoding="utf-8") as f:
        f.write("\n".join(sentences))


def create_csv_file_if_not_exist(file_path: str):
    if not os.path.exists(file_path):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f'{config["pinecone"]["target_column"]}\n')


if __name__ == "__main__":
    main()
