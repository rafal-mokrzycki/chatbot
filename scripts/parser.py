#!/usr/bin/python
import os
import re
from pathlib import Path

import aspose.pdf as pdf
import repackage
import spacy
import win32com.client
from docx import Document
from pdfminer.high_level import extract_text

repackage.up()
from config.config import load_config

config = load_config()


def main(path: str | None = None, verbose: bool = False):
    """
    Main parser of different files. Adds chunks of texts to a local CSV file.

    Args:
        path (str | None, optional): Local path files are in. Defaults to None.
        verbose (bool): Whether to print out info if a file has been parsed.
    """
    list_of_files = get_files_in_dir(path)
    for file_path in list_of_files:
        if file_path.endswith(".pdf"):
            try:
                text = PDFParser(file_path)
                list_of_sentences = text.parse()
                add_lines(list_of_sentences)
                if verbose:
                    print(f"File `{file_path}` parsed")
            except RuntimeError:
                # TODO: problem with ASPOSE client -> check another solution?
                print(f"File `{file_path}` NOT parsed")
        elif file_path.endswith(".docx"):
            text = DOCXParser(file_path)
            list_of_sentences = text.parse()
            add_lines(list_of_sentences)
            if verbose:
                print(f"File `{file_path}` parsed")
        else:
            continue


class DOCXParser:
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
        if "Zarządzenie" in self.file_path:
            self.file_type = "Zarządzenie"
        elif "Pytania" in self.file_path:
            self.file_type = "Pytania"
        else:
            self.file_type = "Inny"

    def parse(self) -> list[str]:
        """
        Sequentially processes DOCX file.

        Returns:
            list[str]: List of sentences
        """
        if self.file_type.capitalize() == "Zarządzenie":
            t1 = extract_text_from_docx(self.file_path)
            t2 = remove_preambule_before_par(t1)
            t3 = remove_attachments(t2)
            t4 = replace_forbidden_chars(t3)
            return split_on_points(t4)
        elif self.file_type.capitalize() == "Pytania":
            t1 = extract_text_from_docx(self.file_path)
            t2 = remove_preambule_before_point(t1)
            t3 = remove_page_numbers(t2)
            t4 = replace_whitespaces(t3)
            t5 = replace_forbidden_chars(t4)
            t6 = add_category(t5, self.file_path)
            return [t6]
        else:
            t1 = extract_text_from_docx(self.file_path)
            t2 = chunk_text(t1)
            t3 = replace_whitespaces(t2)
            return replace_forbidden_chars(t3)


class PDFParser:
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
        if "Sylabus" in self.file_path:
            self.file_type = "Sylabus"
        else:
            self.file_type = "Inny"

    def parse(self):
        if self.file_type.capitalize() == "Sylabus":
            # TODO: get rest of text after tables
            t1 = extract_text_from_pdf(self.file_path)
            header = get_header(text=t1, file_path=self.file_path)
            json_ = jsonize_pdf(t1)
            text_list = prettify_json(json=json_, header=header)
            t2 = replace_forbidden_chars(text_list)
            return replace_whitespaces(t2)
        else:
            # TODO: modify if new PDF files arrive
            t1 = extract_text_from_textual_pdf(self.file_path)
            t2 = replace_forbidden_chars(t1)
            t3 = replace_whitespaces(t2)
            return [t3]


def get_files_in_dir(path_to_search: str | None = None) -> list:
    """
    Searches all files in a dir (checking also subdirs) that have extensions as stated
    in config.json.

    Args:
        path_to_search (str | None, optional): Path to search files in. If None,
        equals to self.path. Defaults to None.

    Returns:
        list: List of files with absolute paths.
    """
    if path_to_search is None:
        path_to_search = Path(__file__).parent.parent.joinpath("data")
    if os.path.isfile(path_to_search):
        return [path_to_search]
    elif os.path.isdir(path_to_search):
        file_paths = []
        for root, _, files in os.walk(path_to_search):
            for file in files:
                _, ext = os.path.splitext(file)
                if ext.lower() in config["general"]["accepted_file_formats"]:
                    file_path = os.path.join(root, file)
                    file_paths.append(os.path.abspath(file_path))
        return file_paths


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts sentences from DOCX file and returns as one string. Joins single and
    multiple '\n' into single whitespaces.

    Args:
        file_path (str): DOCX file path to extract text from.

    Returns:s
        str: Extracted text.
    """
    if file_path.endswith(".docx"):
        doc = Document(file_path)
        text = " ".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    elif file_path.endswith(".doc"):
        word = win32com.client.Dispatch("Word.Application")
        word.visible = False
        wb = word.Documents.Open(file_path)
        doc = word.ActiveDocument
        text = doc.Range().Text
        return text.replace("\r", " ").replace("\n", " ")
    else:
        raise TypeError("File must be DOCX or DOC.")


def split_on_points(text: str) -> list[str]:
    """
    Splits text of paragraphs and points (eg. § 2, §2, $ 2, $2, 2., 13.).

    Args:
        text (str): Text to be divided.

    Returns:
        list[str]: List of remained sentences.
    """
    # dzielenie na paragrafach i punktach TODO: czy dzielimy na podpunktach?
    text_without_paragrapghs = re.split(r"[§$] ?\d+.?", text)
    text_without_points = re.split(r"\d+\.", " ".join(text_without_paragrapghs))
    # usuwanie białych znaków na początku i na końcu stringa
    # oraz usuwanie pustych elementów listy
    sentences = [sentence.strip() for sentence in text_without_points if sentence.strip()]
    # usuwanie wielokrotnych spacji wewnątrz stringów i zwrócenie listy
    return [replace_whitespaces(sentence.strip()) for sentence in sentences]


def remove_preambule_before_par(text: str) -> str:
    """
    Removes everything before first paragraph (§) including this paragraph sign with a
    number. Strips whitespace at the beginnin of the remained text. As first paragraph
    (§) forms may vary, the following substrings are tried to be splitters: "§ 1",
    "§1", "$ 1", "$1". WARNING: this set may not be a finite one.

    Args:
        text (str): Text to be divided.

    Returns:
        str: Remained text body.
    """
    # usuwanie wstępu (przed paragrafem 1.)
    try:
        return re.split(r"[§$] ?1.?", text, 1)[1]
    except IndexError:
        return replace_whitespaces(text.strip())


def remove_preambule_before_point(text: str) -> str:
    """
    Removes everything before first point (1.).

    Args:
        text (str): Text to be divided.

    Returns:
        str: Remained text body.
    """
    # usuwanie wstępu (przed punktem 1.)
    try:
        return f"1. {re.split(r' 1.? ', text, 1)[1]}"
    except IndexError:
        return replace_whitespaces(text.strip())


def remove_attachments(text: str) -> str:
    """
    Removes everything before first attachment (Załącznik nr 1).

    Args:
        text (str): Text to be divided.

    Returns:
        str: Remained text body.
    """
    # usuwanie załączników wraz z tytułem sekcji (np. "Załącznik nr X")
    return replace_whitespaces(re.split(r"Za[lł][aą]cznik ?(nr)? ?1", text, 1)[0]).strip()


def replace_forbidden_chars(
    text: str | list[str],
    replacement_dict: dict | None = None,
) -> str | list[str]:
    """
    Replaces characters that prevent from reading CSV file into pandas DataFrame.

    Args:
        text (str | list[str]): Text or list of texts to apply replacement on.
        replacement_dict (dict): Dictionary with characters to be replaced (keys) and
        characters to replace with (values). If None, taken from config.json.
        Defaults to None.
    Returns:
        (str | list[str]): Text or list of texts.
    """
    if replacement_dict is None:
        replacement_dict = config["general"]["replacement_dictionary"]
    elif not isinstance(replacement_dict, dict):
        raise TypeError("Must be a dictionary.")
    if isinstance(text, str):
        for key in replacement_dict:
            text = text.replace(key, replacement_dict[key])
        return replace_whitespaces(text.strip())
    elif isinstance(text, list):
        result = []
        for elem in text:
            if not isinstance(elem, str):
                raise TypeError("Text must be string or list of strings.")
            for key in replacement_dict:
                elem = elem.replace(key, replacement_dict[key])
            result.append(replace_whitespaces(elem.strip()))
        return result
    else:
        raise TypeError("Text must be string or list of strings.")


def replace_whitespaces(text: str | list[str]) -> str | list[str]:
    """
    Replaces multiple whitespaces with single ones.

    Args:
        text (str | list[str]): Text of list of texts to apply replacement on.

    Returns:
        str | list[str]: Text or list of texts.
    """
    if isinstance(text, str):
        return re.sub(r"\s+", " ", text)
    elif isinstance(text, list):
        return [re.sub(r"\s+", " ", t) for t in text]
    else:
        raise TypeError("Text must be string of list of strings.")


def remove_page_numbers(text: str) -> str:
    """
    Replaces page numbers (i.e. '2 z 5' or '13 z 134').

    Args:
        text (str): Text to apply removal on.

    Returns:
        str: Text.
    """
    return re.sub(r"\d{1,3} z \d{1,3}", "", text)


def add_category(text: str, file_path: str) -> str:
    """
    Returns a formatted string with category retrieved from file name and text.

    Args:
        text (str): Text to operate on.
        file_path (str): File path to modify.

    Returns:
        str: Formatted text.
    """
    category = file_path.split("\\")[-1].split(".")[0]
    result = f"{category}: {text}\n"
    return replace_whitespaces(result.strip())


def determine_language(text: str) -> str:
    """
    Returns a language, that given text is written in.

    Args:
        text (str): Text to check.

    Returns:
        str: Language. Tested for pl, en, de.
    """
    import spacy_fastlang

    nlp = spacy.load("pl_core_news_sm")
    nlp.add_pipe("language_detector")
    doc = nlp(text)

    return doc._.language


def chunk_text(text: str) -> list[str]:
    """
    Chunks text into overlapping 3-sentence chunks (they overlap 1 sentence with previous
    chunk, 1 with next and 1 sentence is non-overlapping.)

    Args:
        text (str): String to chunk.

    Returns:
        list[str]: List of strings after chunking.
    """
    nlp = spacy.load("pl_core_news_sm")
    doc = nlp(text)
    sentences = [sent.text.strip() for sent in doc.sents]
    result = []
    for i, _ in enumerate(sentences):
        if i < len(sentences) - 2:
            pattern = " ".join(str(sentences[j]) for j in range(i, i + 3))
            result.append(replace_whitespaces(pattern.strip()))

    return result


def extract_text_from_pdf(file_path: str) -> str:
    """
    Returns raw text (single string) from a PDF file.

    Args:
        file_path (str): File path to read PDF file from.

    Returns:
        str: Full text of a PDF file
    """
    pdfDocument = pdf.Document(file_path)
    full_text = []
    # TODO: add check for pages number and whether all pages are used in parsing
    for i in range(0, len(pdfDocument.pages)):
        full_text.append(get_raw_text_from_tables(pdfDocument, i + 1))
    result = "".join(full_text)
    return replace_whitespaces(result.strip())


def extract_text_from_textual_pdf(file_path: str) -> str:
    """
    Returns plain text from a PDF file (without tables)

    Args:
        file_path (str): Path to file to extract text from.

    Returns:
        str: One-line text.
    """
    raw_text = extract_text(file_path)
    result = raw_text.replace("\n", " ")
    return replace_whitespaces(result.strip())


def get_raw_text_from_tables(pdfDocument: pdf.Document, iterator: int) -> str:
    """
    Returns raw text (single string) from a table in a PDF file.

    Args:
        pdfDocument (pdf.Document): PDF pages.
        iterator (int): Number of page.

    Returns:
        str: Full text of a PDF file page.
    """
    # TODO: add exception handling if there is no table in the file.
    # TODO: add exception handling if there is a problem with pages number.
    # Initialize TableAbsorber object
    tableAbsorber = pdf.text.TableAbsorber()

    # Parse all the tables on first page
    tableAbsorber.visit(pdfDocument.pages[iterator])

    # Get a reference of the first table
    absorbedTable = tableAbsorber.table_list[0]

    full_text = []

    # Iterate through all the rows in the table
    for pdfTableRow in absorbedTable.row_list:
        # Iterate through all the columns in the row
        for pdfTableCell in pdfTableRow.cell_list:
            # Fetch the text fragments
            textFragmentCollection = pdfTableCell.text_fragments
            # Iterate through the text fragments
            for textFragment in textFragmentCollection:
                # Print the text
                full_text.extend(textFragment.text)
    result = "".join(full_text)
    return replace_whitespaces(result.strip())


def get_header(text: str, file_path: str) -> str:
    """
    Gets header of a Syllabus.

    Args:
        text (str): A text get header from.

    Returns:
        str: A header.
    """
    # Adding ? after the quantifier makes it perform the match in non-greedy or \
    # minimal fashion; as few characters as possible will be matched.
    pattern = r".*sylabus.*? (?=1\.)"
    try:
        header = re.search(pattern=pattern, string=text, flags=re.IGNORECASE)[0]
    except TypeError:
        if "FiR" in file_path:
            header = "Sylabus praktyk zawodowych na kierunku Finanse i Rachunkowość"
        elif "Z" in file_path:
            header = "Sylabus praktyk zawodowych na kierunku Zarządzanie"
        elif "Ekonomia" in file_path:
            header = "Sylabus praktyk zawodowych na kierunku Ekonomia"
        else:
            header = "Sylabus"
    return replace_whitespaces(header.strip())


def jsonize_pdf(text: str) -> dict:
    """
    Gets sought after elements and returns them as a dict, eg.
    dict['Metody kształcenia'] = 'Wskazane przez praktykodawcę'

    Args:
        main_string (str): String to search in.

    Returns:
        dict: Dictionary with `headers` as keys and information as values.
    """
    json_ = {}
    list_of_headers = [
        "Nazwa przedmiotu",
        "Forma zajęć",
        "Rok akademicki, rok studiów, semestr realizacji przedmiotu",
        "Stopień studiów, tryb studiów",
        "Cel przedmiotu",
        "Wymagania wstępne",
        "Metody kształcenia",
    ]
    for elem in list_of_headers:
        # result = forward_regexp_search(text, elem)
        try:
            result = re.search(rf"{elem}(.*?)\d+\.", text).group(1)
            json_[elem] = replace_whitespaces(result.strip())
        except AttributeError:
            pass
    return json_


def remove_keywords(text: str, to_remove: list | None = None) -> str:
    """
    Removes certain keywords from string.

    Args:
        text (str): String to operate on.
        to_remove (list | None): Strings to remove. Defaults to None.

    Returns:
        str: New string.
    """
    if to_remove is None:
        to_remove = [
            "Wiedza",
            "Kod efektu",
            "Metody weryfikacji",
            "Przedmiotowy",
            "Kierunkowy",
        ]
    for token in to_remove:
        text = text.replace(token, "")
    return replace_whitespaces(text.strip())


def remove_codes(text: str) -> str:
    """
    Removes certain codes ('kody efektów kształcenia') from a string based on regexp.

    Args:
        text (str): String to operate on.

    Returns:
        str: New string.
    """
    pattern = r"(EP-\d{1,2})|(K_[KUW]?\d{1,2})"
    result = re.sub(pattern=pattern, repl="", string=text)
    return replace_whitespaces(result.strip())


def forward_regexp_search(main_string: str, target_string: str) -> str:
    """
    Searches a given substring (`target_string`) in a large string (`main_string`)
    and returns N next words as one string (up to the nearest number with a comma, eg.
    `13.`)

    Args:
        main_string (str): String to search in
        target_string (str): Subtring to search

    Returns:
        str: Searched string.
    """

    def get_number_with_comma(string):
        regexp = r"\d+\."
        if re.fullmatch(regexp, string) is not None:
            return False
        return True

    # get rid of whitespaces at the beginning and in the end
    target_string = target_string.strip()
    length = len(target_string)
    # position of the target_string
    pos = main_string.find(target_string)
    word_list = (main_string[pos + length :].strip()).split(" ")
    result_lst = []
    for word in word_list:
        if get_number_with_comma(word):
            result_lst.append(word)
        else:
            break
    result = " ".join(result_lst)
    return replace_whitespaces(result.strip())


def prettify_json(json: dict, header: str) -> list[str]:
    """
    Prettifies a dictionary. Changes {"Nazwa przedmiotu" : "praktyka"} to
    ["Nazwa przedmiotu na kierunku 'kierunek': praktyka"]

    Args:
        json (dict): Python dictionary to change.
        header (str): Header to extract faculty from.

    Returns:
        list[str]: Formatted string.
    """
    new_json = dict()
    if "Zarządzanie" in header:
        faculty_name = "Zarządzanie"
    elif "Finanse i Rachunkowość" in header:
        faculty_name = "Finanse i Rachunkowość"
    else:
        faculty_name = "Ekonomia"
    for key in json:
        new_key = f"{key} przedmiotu praktyka zawodowa na kierunku {faculty_name}: "
        new_json[new_key] = json[key]
    result = []
    for key in new_json:
        result.append(f"{key}{new_json[key]}")
    return result


def add_lines(sentences: list[str], file_path: str | None = None):
    """
    Appends sentences to file. Creates a file if it does no exist.

    Args:
        sentences (list[str]): Sentences to append.
        file_path (str | None, optional): CSV file path to append sentences to.
        If None, taken from config.json. Defaults to None.
    """
    if file_path is None:
        file_name = config["pinecone"]["target_filename"]["raw"]
        file_path = Path(__file__).parent.parent.joinpath(file_name)
    create_csv_file_if_not_exist(file_path)
    with open(file_path, "a", encoding="utf-8") as f:
        f.write("\n".join(sentences))


def create_csv_file_if_not_exist(file_path: str):
    """
    Creates a CSV file of a given name if file does not exist.

    Args:
        file_path (str): Full path to create CSV file in.
    """
    if not os.path.exists(file_path):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f'{config["pinecone"]["target_column"]}\n')


if __name__ == "__main__":
    main(verbose=False)
